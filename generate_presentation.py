from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# Create document
doc = Document()

# Set styles
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ==================== TITLE PAGE ====================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run('RISK DETECTOR AI')
title_run.font.size = Pt(36)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(61, 220, 132)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run('Métricas de Detecção de Jogadores de Risco')
subtitle_run.font.size = Pt(18)
subtitle_run.font.color.rgb = RGBColor(160, 160, 160)

doc.add_paragraph()

# Date
date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_run = date_para.add_run(f'Apresentação - {datetime.now().strftime("%d de %B de %Y").replace("February", "Fevereiro")}')
date_run.font.size = Pt(12)
date_run.font.italic = True

# ==================== SECTION 1: OVERVIEW ====================
doc.add_page_break()

heading1 = doc.add_heading('1. Visão Geral do Modelo', level=1)
heading1.runs[0].font.color.rgb = RGBColor(61, 220, 132)

doc.add_paragraph(
    'O Risk Detector AI é um sistema de detecção de anomalias que identifica jogadores '
    'com comportamento de risco elevado em plataformas de apostas esportivas. O modelo utiliza '
    'algoritmo de Isolation Forest para análise não-supervisionada.',
    style='Normal'
)

# Algoritmo
doc.add_heading('Algoritmo Utilizado', level=2)
p = doc.add_paragraph()
p.add_run('Isolation Forest').bold = True
p.add_run(' - Detecção de anomalias baseada em floresta de decisão')

doc.add_paragraph('Aprende o padrão "normal" de comportamento dos usuários', style='List Bullet')
doc.add_paragraph('Identifica desvios estatísticos significativos', style='List Bullet')
doc.add_paragraph('Não requer dados rotulados (aprendizado não-supervisionado)', style='List Bullet')
doc.add_paragraph('Processamento rápido e escalável', style='List Bullet')

# Configuração
doc.add_heading('Configuração Atual', level=2)
table = doc.add_table(rows=5, cols=2)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Parâmetro'
table.rows[0].cells[1].text = 'Valor'
table.rows[1].cells[0].text = 'N° de Estimadores'
table.rows[1].cells[1].text = '300'
table.rows[2].cells[0].text = 'Contaminação (Anomalias)'
table.rows[2].cells[1].text = '1% (0.01)'
table.rows[3].cells[0].text = 'Random State'
table.rows[3].cells[1].text = '42'
table.rows[4].cells[0].text = 'Processamento Paralelo'
table.rows[4].cells[1].text = 'Ativado'

# ==================== SECTION 2: METRICS ====================
doc.add_page_break()

heading2 = doc.add_heading('2. Métricas de Análise', level=1)
heading2.runs[0].font.color.rgb = RGBColor(61, 220, 132)

doc.add_paragraph(
    'O modelo extrai 6 métricas agregadas por usuário que capturam padrões de risco '
    'em diferentes dimensões do comportamento de apostas.',
    style='Normal'
)

# Metric 1
doc.add_heading('2.1 Contagem de Apostas (bet_count)', level=2)
doc.add_paragraph(
    'Total de apostas realizadas pelo usuário no período analisado.',
    style='Normal'
)
doc.add_paragraph('O quê identifica:', style='List Bullet')
doc.add_paragraph('Usuários com atividade frequente de apostas', style='List Bullet 2')
doc.add_paragraph('Padrões de engajamento intenso', style='List Bullet 2')
doc.add_paragraph('Comportamento compulsivo potencial', style='List Bullet 2')
doc.add_paragraph('Intervalo esperado: 10 a 500+ apostas por usuário', style='Normal')

# Metric 2
doc.add_heading('2.2 Valor Total Apostado (total_stake)', level=2)
doc.add_paragraph(
    'Somatório de todos os valores investidos em apostas pelo usuário.',
    style='Normal'
)
doc.add_paragraph('O quê identifica:', style='List Bullet')
doc.add_paragraph('Volume financeiro investido (exposição)', style='List Bullet 2')
doc.add_paragraph('Usuários com apostas de alto valor', style='List Bullet 2')
doc.add_paragraph('Padrão de risco financeiro', style='List Bullet 2')
doc.add_paragraph('Intervalo esperado: R$ 50 a R$ 50.000+', style='Normal')

# Metric 3
doc.add_heading('2.3 Valor Médio por Aposta (avg_stake)', level=2)
doc.add_paragraph(
    'Média dos valores investidos por aposta (total_stake / bet_count).',
    style='Normal'
)
doc.add_paragraph('O quê identifica:', style='List Bullet')
doc.add_paragraph('Tamanho típico das apostas do usuário', style='List Bullet 2')
doc.add_paragraph('Perfil de aversão/busca ao risco', style='List Bullet 2')
doc.add_paragraph('Padrão de comportamento consistente', style='List Bullet 2')
doc.add_paragraph('Intervalo esperado: R$ 5 a R$ 500 por aposta', style='Normal')

# Metric 4
doc.add_heading('2.4 Odds Médias (avg_odds)', level=2)
doc.add_paragraph(
    'Média das odds (probabilidades inversas) das apostas realizadas.',
    style='Normal'
)
doc.add_paragraph('O quê identifica:', style='List Bullet')
doc.add_paragraph('Nível de risco matemático das escolhas', style='List Bullet 2')
doc.add_paragraph('Odds altas = maior risco e potencial de perda', style='List Bullet 2')
doc.add_paragraph('Padrão de busca por retornos altos', style='List Bullet 2')
doc.add_paragraph('Intervalo esperado: 1.5 a 10.0 (quanto maior = mais risco)', style='Normal')

# Metric 5
doc.add_heading('2.5 Taxa de Retorno (return_ratio)', level=2)
doc.add_paragraph(
    'Razão entre ganhos totais e investimento total (gain_amount / total_stake).',
    style='Normal'
)
doc.add_paragraph('O quê identifica:', style='List Bullet')
doc.add_paragraph('Performance de apostas (lucratividade)', style='List Bullet 2')
doc.add_paragraph('Taxa de sucesso (valores > 1.0 = lucro)', style='List Bullet 2')
doc.add_paragraph('Padrão de comportamento de perdedor', style='List Bullet 2')
doc.add_paragraph('Valores negativos indicam perdas sistemáticas', style='List Bullet 2')
doc.add_paragraph('Intervalo esperado: -2.0 a 5.0', style='Normal')

# Metric 6
doc.add_heading('2.6 Dias desde Última Aposta (days_since_last_bet)', level=2)
doc.add_paragraph(
    'Número de dias entre a última aposta registrada e a data de análise.',
    style='Normal'
)
doc.add_paragraph('O quê identifica:', style='List Bullet')
doc.add_paragraph('Recência de atividade (usuários ativos vs inativos)', style='List Bullet 2')
doc.add_paragraph('Padrões de ausência seguida de retorno', style='List Bullet 2')
doc.add_paragraph('Picos de atividade irregular', style='List Bullet 2')
doc.add_paragraph('Intervalo esperado: 0 a 365+ dias', style='Normal')

# ==================== SECTION 3: RISK DETECTION ====================
doc.add_page_break()

heading3 = doc.add_heading('3. Como o Modelo Detecta Risco', level=1)
heading3.runs[0].font.color.rgb = RGBColor(61, 220, 132)

doc.add_paragraph(
    'O modelo identifica ANOMALIAS - combinações de métricas que se desviam do padrão normal. '
    'Usuários são considerados de risco quando apresentam comportamentos estatisticamente incomuns.',
    style='Normal'
)

doc.add_heading('Exemplos de Padrões Flagged', level=2)

patterns = [
    ('Alto Engajamento + Alto Investimento', 'Muitas apostas de valores altos consecutivos'),
    ('Odds Extremas', 'Apostas sistemáticas com odds muito altas (> 7.0)'),
    ('Taxa de Retorno Negativa Extrema', 'Perdas sistemáticas > 50% do investimento'),
    ('Picos de Atividade', 'Surtos de apostas em curtos períodos seguidos de inatividade'),
    ('Combinação Anômala', 'Padrão único que não segue o perfil típico dos usuários'),
]

for pattern, description in patterns:
    p = doc.add_paragraph()
    p.add_run(f'{pattern}: ').bold = True
    p.add_run(description)
    p.style = 'List Bullet'

doc.add_heading('Pré-processamento de Dados', level=2)
doc.add_paragraph('Antes da análise, os dados são normalizados:', style='Normal')
doc.add_paragraph('Imputação de valores faltantes (mediana)', style='List Bullet')
doc.add_paragraph('Normalização com StandardScaler (média=0, desvio=1)', style='List Bullet')
doc.add_paragraph('Remoção de outliers extremos', style='List Bullet')
doc.add_paragraph('Conversão de todas as métricas para escala comparável', style='List Bullet')

# ==================== SECTION 4: SCORE INTERPRETATION ====================
doc.add_page_break()

heading4 = doc.add_heading('4. Interpretação dos Scores de Risco', level=1)
heading4.runs[0].font.color.rgb = RGBColor(61, 220, 132)

doc.add_paragraph(
    'Cada usuário recebe um score de risco normalizado entre 0.0 (baixo risco) e 1.0 (risco máximo).',
    style='Normal'
)

# Score table
risk_table = doc.add_table(rows=6, cols=3)
risk_table.style = 'Light Grid Accent 1'
risk_table.rows[0].cells[0].text = 'Faixa de Score'
risk_table.rows[0].cells[1].text = 'Nível'
risk_table.rows[0].cells[2].text = 'Significado'

risk_table.rows[1].cells[0].text = '0.00 - 0.39'
risk_table.rows[1].cells[1].text = '✅ Baixo'
risk_table.rows[1].cells[2].text = 'Comportamento dentro do padrão esperado'

risk_table.rows[2].cells[0].text = '0.40 - 0.59'
risk_table.rows[2].cells[1].text = '⚡ Médio'
risk_table.rows[2].cells[2].text = 'Alguns indicadores de desvio, requer acompanhamento'

risk_table.rows[3].cells[0].text = '0.60 - 0.79'
risk_table.rows[3].cells[1].text = '⚠️ Alto'
risk_table.rows[3].cells[2].text = 'Comportamento claramente anômalo, ação recomendada'

risk_table.rows[4].cells[0].text = '0.80 - 1.00'
risk_table.rows[4].cells[1].text = '🚨 Crítico'
risk_table.rows[4].cells[2].text = 'Risco extremo, intervenção imediata necessária'

doc.add_heading('Ações Recomendadas por Nível', level=2)

actions = {
    'Baixo (< 0.4)': ['Monitoramento padrão', 'Não requer ação imediata'],
    'Médio (0.4 - 0.59)': ['Acompanhamento regular', 'Revisar padrão de apostas', 'Verificar mudanças de comportamento'],
    'Alto (0.6 - 0.79)': ['Contato com usuário', 'Análise detalhada de histórico', 'Possível limitação de apostas', 'Oferecer ajuda responsável'],
    'Crítico (≥ 0.8)': ['Ação imediata obrigatória', 'Bloqueio temporário de conta', 'Contato urgente', 'Possível encaminhamento para proteção'],
}

for level, action_list in actions.items():
    p = doc.add_paragraph()
    p.add_run(f'{level}:').bold = True
    p.style = 'List Bullet'
    for action in action_list:
        sub_p = doc.add_paragraph(action, style='List Bullet 2')

# ==================== SECTION 5: DATA TRAINING ====================
doc.add_page_break()

heading5 = doc.add_heading('5. Dados de Treinamento', level=1)
heading5.runs[0].font.color.rgb = RGBColor(61, 220, 132)

train_table = doc.add_table(rows=5, cols=2)
train_table.style = 'Light Grid Accent 1'
train_table.rows[0].cells[0].text = 'Métrica'
train_table.rows[0].cells[1].text = 'Valor'
train_table.rows[1].cells[0].text = 'Total de Registros de Apostas'
train_table.rows[1].cells[1].text = '22.844'
train_table.rows[2].cells[0].text = 'Usuários Únicos Identificados'
train_table.rows[2].cells[1].text = '52'
train_table.rows[3].cells[0].text = 'Período de Análise'
train_table.rows[3].cells[1].text = 'Histórico completo'
train_table.rows[4].cells[0].text = 'Features Selecionadas'
train_table.rows[4].cells[1].text = '4 numéricas'

doc.add_paragraph()

doc.add_heading('Contaminação Esperada', level=2)
doc.add_paragraph(
    f'Com contamination=0.01 (1%), o modelo espera identificar aproximadamente 1% dos usuários '
    f'como anomalias significativas. Isso garante alta assertividade nas detecções, '
    f'reduzindo falsos positivos e focando em casos verdadeiramente preocupantes.',
    style='Normal'
)

# ==================== SECTION 6: ADVANTAGES ====================
doc.add_page_break()

heading6 = doc.add_heading('6. Vantagens do Modelo', level=1)
heading6.runs[0].font.color.rgb = RGBColor(61, 220, 132)

advantages = [
    'Não supervisionado - não requer dados rotulados',
    'Detecta novos padrões de risco não conhecidos a priori',
    'Altamente escalável e rápido em processamento',
    'Interpretável - baseado em métricas de negócio reais',
    'Adaptável - pode ser retreinado com novos dados',
    'Reduz carga de revisão manual com detecções assertivas',
    'Integração simples com sistemas existentes',
    'Dashboard intuitivo para análise visual',
]

for adv in advantages:
    doc.add_paragraph(adv, style='List Bullet')

# ==================== SECTION 7: LIMITATIONS ====================
doc.add_heading('7. Limitações e Considerações', level=1)
heading7 = doc.add_heading('7. Limitações e Considerações', level=1)
heading7.runs[0].font.color.rgb = RGBColor(255, 107, 107)

limitations = [
    'Depende da qualidade e completude dos dados históricos',
    'Anomalias legítimas podem ser confundidas com risco',
    'Requer validação manual de casos críticos',
    'Performance pode variar com mudanças no padrão de mercado',
    'Não substitui análise humana especializada',
]

for lim in limitations:
    doc.add_paragraph(lim, style='List Bullet')

# ==================== FOOTER ====================
doc.add_page_break()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer.add_run('---')
footer_run.font.color.rgb = RGBColor(61, 220, 132)

final = doc.add_paragraph()
final.alignment = WD_ALIGN_PARAGRAPH.CENTER
final_run = final.add_run('Documento preparado para apresentação\nRisk Detector AI - Detecção de Jogadores de Risco')
final_run.font.size = Pt(10)
final_run.font.italic = True
final_run.font.color.rgb = RGBColor(160, 160, 160)

# Save
output_path = r'c:\Users\Caio Araujo\Documents\Programação\Programas Esportiva\risk_detector_ai\METRICAS_RISK_DETECTOR.docx'
doc.save(output_path)
print(f'✅ Documento criado com sucesso: {output_path}')
